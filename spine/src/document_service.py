from docx import Document
from spine.src.models import ClauseNode
import os

import re

class DocumentParser:
    def load(self, filepath: str) -> ClauseNode:
        if not os.path.exists(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")

        doc = Document(filepath)
        return self._parse_document(doc)

    def parse_stream(self, stream) -> ClauseNode:
        """Parse a document from a file-like object (bytes)"""
        doc = Document(stream)
        return self._parse_document(doc)

    def _parse_document(self, doc: Document) -> ClauseNode:
        root = ClauseNode(an_type="document", text="ROOT")
        
        current_article = None
        current_section = None
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            style_name = p.style.name
            
            # Safe ID strategy
            try:
                if p._element is not None:
                    para_id = p._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
                    if not para_id:
                        para_id = str(id(p._element))
                else:
                    para_id = str(id(p))
            except Exception:
                para_id = str(uuid.uuid4())

            # --- Detection Logic ---
            
            # 1. Regex Detection for Numbering
            # e.g. "1.1 Definitions" -> num="1.1", type=section/article
            # e.g. "(a) Item" -> num="(a)", type=point
            
            # Patterns
            pattern_section = r'^(\d+(\.\d+)*)\.?\s+(.*)' # 1.1 or 1.1.1 or 4
            pattern_point = r'^(\([a-z0-9]+\))\s+(.*)' # (a) or (1)
            
            node = None
            
            # Priority 1: Headers (Explicit)
            if style_name.startswith('Heading 1') or style_name == 'Title':
                node = ClauseNode(an_type="article", text=text, original_xml_id=para_id)
                # Try to extract num from text if present, e.g. "Article 1: Foo"
                # For now assume headings are structural enough.
                root.children.append(node)
                current_article = node
                current_section = None 
                
            elif style_name.startswith('Heading 2'):
                # Try to parse num
                match = re.match(pattern_section, text)
                an_num = match.group(1) if match else None
                
                node = ClauseNode(an_type="section", text=text, original_xml_id=para_id, an_num=an_num)
                if current_article:
                    current_article.children.append(node)
                else:
                    root.children.append(node)
                current_section = node
            
            # Priority 2: Regex on Normal/List Text
            else:
                s_match = re.match(pattern_section, text)
                p_match = re.match(pattern_point, text)
                
                if s_match:
                     # It looks like a section (1.1 Foo) but is styled as Normal. Treat as Section.
                     an_num = s_match.group(1)
                     node = ClauseNode(an_type="section", text=text, original_xml_id=para_id, an_num=an_num)
                     if current_article:
                         current_article.children.append(node)
                     else:
                         root.children.append(node)
                     current_section = node
                     
                elif p_match:
                    # It looks like a point (a) Foo
                    an_num = p_match.group(1)
                    node = ClauseNode(an_type="point", text=text, original_xml_id=para_id, an_num=an_num)
                    if current_section:
                        current_section.children.append(node)
                    elif current_article:
                         current_article.children.append(node)
                    else:
                        root.children.append(node)
                        
                else:
                    # Standard Paragraph
                    node = ClauseNode(an_type="paragraph", text=text, original_xml_id=para_id)
                    if current_section:
                        current_section.children.append(node)
                    elif current_article:
                        current_article.children.append(node)
                    else:
                        root.children.append(node)
                    
        return root

    def print_tree(self, node: ClauseNode, level=0):
        indent = "  " * level
        print(f"{indent}[{node.an_type}] {node.text[:50]}...")
        for child in node.children:
            self.print_tree(child, level + 1)
