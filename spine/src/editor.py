from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from spine.src.models import ClauseNode
import copy

class DocumentEditor:
    def __init__(self, original_docx_path: str):
        self.doc = Document(original_docx_path)
        # Build a map of xml_id -> paragraph element
        self.xml_map = {}
        for p in self.doc.paragraphs:
            # Use same strategy: w:paraId or id(element)
            para_id = p._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
            if not para_id:
                para_id = str(id(p._element))
            self.xml_map[para_id] = p

    def save(self, output_path: str):
        self.doc.save(output_path)

    def inject_paragraph_after(self, target_node_id_xml: str, new_text: str, style="Normal"):
        """
        Injects a new paragraph after the paragraph with target_node_id_xml.
        """
        if target_node_id_xml not in self.xml_map:
            raise ValueError(f"Target XML ID {target_node_id_xml} not found in document.")

        target_p = self.xml_map[target_node_id_xml]
        
        # Create new paragraph element
        new_p = self.doc.add_paragraph(new_text, style=style)
        # Remve it from the end where add_paragraph put it
        # self.doc._body._element.remove(new_p._element) # This is tricky in python-docx
        # Easier: Create a new paragraph element manually using OxmlElement if we want clean control, 
        # but using the high level API and moving the element is safer for styles.
        
        # Strategy: The 'new_p' is at the bottom of the document body. 
        # We need to move 'new_p._element' to be after 'target_p._element'.
        
        body = self.doc._body._element
        new_p_element = new_p._element
        
        # Detach from bottom
        body.remove(new_p_element)
        
        # Insert after target
        target_element = target_p._element
        target_index = body.index(target_element)
        body.insert(target_index + 1, new_p_element)
        
        # Return ID of new element
        new_para_id = new_p_element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
        if not new_para_id:
            new_para_id = str(id(new_p_element))
        
        return new_para_id

    def split_paragraph(self, target_node_id_xml: str, split_text_part_a: str, split_text_part_b: str, type_b="List Paragraph"):
        """
        Splits a paragraph into two.
        Strategy: Update target with Part A. inject new paragraph with Part B after it.
        Advanced: Handles numbering for Part B (mocked as manual text prefix for now).
        """
        if target_node_id_xml not in self.xml_map:
            raise ValueError(f"Target XML ID {target_node_id_xml} not found.")

        target_p = self.xml_map[target_node_id_xml]
        
        # Update text of A
        target_p.text = split_text_part_a
        
        # Inject B
        # In Phase 2.5, we often want the split part to become a child item or same level.
        # For the "Refactor Split" use case (4.2 -> 4.2.1, 4.2.2), we actually need to change A's style too potentially.
        # For this spike, we just inject B.
        
        # Inherit style from A or use override
        style = type_b if type_b else target_p.style
        
        # Inject
        self.inject_paragraph_after(target_node_id_xml, split_text_part_b, style=style)
