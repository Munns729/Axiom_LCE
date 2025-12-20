import pytest
import os
from docx import Document
from spine.src.editor import DocumentEditor
from spine.src.document_service import DocumentParser

CORPUS_DIR = os.path.join(os.path.dirname(__file__), 'corpus')

@pytest.fixture
def messy_2_path(tmp_path):
    # Copy corpus file to tmp_path to avoid modifying original
    import shutil
    original = os.path.join(CORPUS_DIR, 'messy_2.docx')
    target = tmp_path / "messy_2_test.docx"
    shutil.copy(original, target)
    return str(target)

def test_injection(messy_2_path):
    editor = DocumentEditor(messy_2_path)
    
    # Find "1.1. The Landlord hereby leases to Tenant."
    # We need to find its XML ID first.
    # In a real app we'd get this from the parsed tree.
    target_id = None
    for p in editor.doc.paragraphs:
        if "1.1. The Landlord" in p.text:
            # Replicate ID strategy
            para_id = p._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
            if not para_id:
                para_id = str(id(p._element))
            target_id = para_id
            break
    
    assert target_id is not None
    
    # Inject
    editor.inject_paragraph_after(target_id, "1.2 INJECTED PARAGRAPH", style="Normal")
    
    output_path = messy_2_path.replace(".docx", "_injected.docx")
    editor.save(output_path)
    
    # Verify
    verify_doc = Document(output_path)
    found = False
    for i, p in enumerate(verify_doc.paragraphs):
        if "1.2 INJECTED PARAGRAPH" in p.text:
            found = True
            # Verify position: Should be after 1.1
            prev_text = verify_doc.paragraphs[i-1].text
            assert "1.1. The Landlord" in prev_text
            break
            
    assert found

def test_split_refactor(messy_2_path):
    editor = DocumentEditor(messy_2_path)
    
    # Find 4.2 Termination
    target_id = None
    for p in editor.doc.paragraphs:
        if "4.2 Termination" in p.text:
            para_id = p._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId')
            if not para_id:
                para_id = str(id(p._element))
            target_id = para_id
            break
            
    assert target_id is not None
    
    # Split
    part_a = "4.2 Termination. This agreement may be terminated by either party."
    part_b = "4.2.1 (Refactored) If the Tenant terminates, they must pay a fee."
    
    editor.split_paragraph(target_id, part_a, part_b)
    
    output_path = messy_2_path.replace(".docx", "_split.docx")
    editor.save(output_path)
    
    # Verify
    verify_doc = Document(output_path)
    found_a = False
    found_b = False
    
    for i, p in enumerate(verify_doc.paragraphs):
        if part_a in p.text:
            found_a = True
            # Next one should be B
            if i + 1 < len(verify_doc.paragraphs):
                 if part_b in verify_doc.paragraphs[i+1].text:
                     found_b = True
    
    assert found_a
    assert found_b
