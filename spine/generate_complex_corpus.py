import os
import random
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CORPUS_DIR = os.path.join(os.path.dirname(__file__), 'tests', 'corpus')
os.makedirs(CORPUS_DIR, exist_ok=True)

def add_p(doc, text, style=None, bold=False, indent=False):
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph(text)
    
    if bold and p.runs:
        p.runs[0].bold = True
        
    if indent:
        p.paragraph_format.left_indent = Pt(36)
        
    # Force stable ID for testing
    hex_id = '{:08X}'.format(random.randint(0, 0xFFFFFFFF))
    p._element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId', hex_id)
    return p

def create_series_a_document():
    doc = Document()
    
    # Title
    title = doc.add_heading('SERIES A INVESTMENT AGREEMENT', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    add_p(doc, 'THIS AGREEMENT is made on this 28th day of January, 2026.')
    
    # SECTION 1: DEFINITIONS
    doc.add_heading('SECTION 1: DEFINITIONS', 1)
    add_p(doc, '1.1 In this Agreement, the following terms shall have the following meanings:')
    
    add_p(doc, '"Bad Leaver" means any Founder who ceases to be an employee or officer of the Company by reason of voluntary resignation or dismissal for Cause.', indent=True)
    add_p(doc, '"Cause" means any material breach of this Agreement, gross negligence, or conviction of a felony.', indent=True)
    add_p(doc, '"Founder" means John Doe and Jane Smith.', indent=True)
    add_p(doc, '"Good Reason" means a material reduction in the Founder\'s duties, base salary, or a relocation of the Company\'s primary office by more than 50 miles without the Founder\'s consent.', indent=True)
    add_p(doc, '"Vesting Period" means the period of four (4) years from the Effective Date.', indent=True)
    
    # SECTION 2: INVESTMENT
    doc.add_heading('SECTION 2: INVESTMENT AND SHARES', 1)
    add_p(doc, '2.1 The Investors agree to purchase the Shares at the Purchase Price on the Completion Date.')
    add_p(doc, '2.2 The Shares shall rank pari passu in all respects with the existing ordinary shares.')
    
    # SECTION 3: VESTING
    doc.add_heading('SECTION 3: VESTING PROVISIONS', 1)
    add_p(doc, '3.1 All Shares held by the Founders shall be subject to reverse vesting over the Vesting Period.')
    add_p(doc, '3.2 Twenty-five percent (25%) of the Shares shall vest on the first anniversary of the Effective Date, with the remainder vesting in equal monthly installments thereafter.')
    
    # SECTION 4: TERMINATION AND FORFEITURE (The Conflict Zone)
    doc.add_heading('SECTION 4: TERMINATION AND FORFEITURE', 1)
    add_p(doc, '4.1 If a Founder ceases to be an employee or officer of the Company, their Unvested Shares shall be subject to compulsory transfer to the Company at par value.')
    
    add_p(doc, '4.2 GOOD REASON PROTECTION. Notwithstanding Section 4.1, if a Founder terminates their employment for Good Reason, they shall be entitled to retain all of their Shares, whether vested or unvested, and no compulsory transfer shall occur.')
    
    add_p(doc, '4.3 VOLUNTARY RESIGNATION OVERRIDE. For the avoidance of doubt, any voluntary resignation by a Founder, regardless of the circumstances or any alleged "Good Reason", shall be deemed a "Bad Leaver" event. In such case, ALL Shares (both vested and unvested) held by the Founder shall be forfeited and transferred to the Company for nil consideration.')
    
    # SECTION 5: WARRANTIES
    doc.add_heading('SECTION 5: WARRANTIES AND LIABILITY', 1)
    add_p(doc, '5.1 The Founders hereby warrant to the Investors that the statements in the Disclosure Letter are true and accurate.')
    add_p(doc, '5.2 LIABILITY CAP. The total aggregate liability of the Founders for any breach of Warranty shall be limited to the Purchase Price.')
    add_p(doc, '5.3 Notwithstanding Section 5.2, in no event shall the Founders\' liability exceed $500,000 in aggregate.')
    
    # SECTION 6: GOVERNING LAW
    doc.add_heading('SECTION 6: GOVERNING LAW', 1)
    add_p(doc, '6.1 This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware.')
    
    file_path = os.path.join(CORPUS_DIR, 'series_a_complex.docx')
    doc.save(file_path)
    return file_path

if __name__ == "__main__":
    path = create_series_a_document()
    print(f"Generated complex Series A document at: {path}")
