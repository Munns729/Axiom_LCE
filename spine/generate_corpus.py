import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

CORPUS_DIR = os.path.join(os.path.dirname(__file__), 'tests', 'corpus')
os.makedirs(CORPUS_DIR, exist_ok=True)

import random

def add_p(doc, text, style=None):
    if style:
        p = doc.add_paragraph(text, style=style)
    else:
        p = doc.add_paragraph(text)
    
    # Force stable ID for testing
    if p._element.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId') is None:
         # Use hash of text for repeatability in this specific test script context, OR random
         # But better to just be random but persistent once saved.
         hex_id = '{:08X}'.format(random.randint(0, 0xFFFFFFFF))
         p._element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId', hex_id)
    return p

def create_messy_doc_1():
    doc = Document()
    doc.add_heading('Messy Contract 1', 0)
    
    # Initial Heading doesn't use helper, might lack ID. But Parser handles headings.
    
    add_p(doc, 'This is a plain paragraph.')
    
    p2 = doc.add_paragraph()
    run = p2.add_run('Section 1.1: Definitions')
    run.bold = True
    run.font.size = Pt(14)
    # inject ID
    hex_id = '{:08X}'.format(random.randint(0, 0xFFFFFFFF))
    p2._element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId', hex_id)
    
    add_p(doc, 'Here is a definition of "Company". It means the company.')

    add_p(doc, '(a) First item', style='List Paragraph')
    add_p(doc, '(b) Second item', style='List Paragraph')

    doc.save(os.path.join(CORPUS_DIR, 'messy_1.docx'))

def create_messy_doc_2():
    doc = Document()
    head = doc.add_heading('LEASE AGREEMENT', 0)
    head.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Helper won't work for complicated run modification, so do manual
    p = doc.add_paragraph('ARTICLE 1: PREMISES')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    p._element.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}paraId', '{:08X}'.format(random.randint(0, 0xFFFFFFFF)))

    add_p(doc, '1.1. The Landlord hereby leases to Tenant.')
    
    text = "4.2 Termination. This agreement may be terminated by either party. If the Tenant terminates, they must pay a fee. If the Landlord terminates, they must provide notice."
    add_p(doc, text)

    doc.save(os.path.join(CORPUS_DIR, 'messy_2.docx'))

if __name__ == "__main__":
    create_messy_doc_1()
    create_messy_doc_2()
    print(f"Generated corpus in {CORPUS_DIR}")
