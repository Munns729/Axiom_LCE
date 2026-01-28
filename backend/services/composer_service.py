from docx import Document
from docx.oxml.ns import qn
import io
from typing import List, Dict, Optional

class ComposerService:
    """
    The 'Bridge' between the Shadow Tree (logic) and the Skeleton DOCX (storage).
    Performs surgical edits based on stable paraIds.
    """

    def __init__(self, file_content: bytes):
        self.source_bytes = file_content
        try:
            self.doc = Document(io.BytesIO(file_content))
            self._build_id_map()
        except Exception as e:
            raise ValueError(f"Failed to load skeleton: {e}")

    def _build_id_map(self):
        """Index all paragraphs by their w:paraId"""
        self.para_map = {}
        for p in self.doc.paragraphs:
            # We assume ID Normalizer has run, so paraId SHOULD exist.
            # But we handle the None case gracefully.
            para_id = p._element.get(qn('w:paraId'))
            if para_id:
                self.para_map[para_id] = p

    def apply_operations(self, operations: List[Dict]) -> bytes:
        """
        Apply a sequence of explicit operations to the document.
        Operations schema:
        [
            {"type": "update_text", "id": "uuid", "text": "New content"},
            # Phase B ops:
            # {"type": "split", "id": "uuid", "parts": ["Part A", "Part B"]}
        ]
        """
        for op in operations:
            op_type = op.get("type")
            node_id = op.get("id")
            
            if not node_id in self.para_map:
                print(f"Warning: Node {node_id} not found in skeleton. Skipping.")
                continue
                
            target_para = self.para_map[node_id]

            if op_type == "update_text":
                self._update_paragraph_text(target_para, op.get("text", ""))
            
            elif op_type == "split":
                # Phase B: The Splitter
                self._split_paragraph(target_para, op.get("parts", []))
        
        return self._save()

    def _split_paragraph(self, target_para, parts: List[str]):
        """
        Split a paragraph into N paragraphs, cloning the style of the original.
        The target_para becomes the first part. New paragraphs are inserted after.
        """
        if not parts:
            return

        # 1. Update the original (first part)
        self._update_paragraph_text(target_para, parts[0])
        
        # 2. Insert subsequent parts
        current_para = target_para
        
        import copy
        import uuid
        
        for i in range(1, len(parts)):
            new_text = parts[i]
            
            # Create a new paragraph by copying the logic of the previous one
            # note: python-docx doesn't have a clean 'clone' method, so we work with XML
            
            # Get the XML element of the current paragraph
            curr_element = current_para._element
            
            # Create a deep copy of the element
            new_element = copy.deepcopy(curr_element)
            
            # Assign a new unique paraId (Critical for future edits)
            # Word expects 8-char hex, but UUID is robust and accepted
            new_id = str(uuid.uuid4())
            new_element.set(qn('w:paraId'), new_id)
            
            # Insert the new element after the current one
            curr_element.addnext(new_element)
            
            # Now wrap it in a proxy object so we can edit it safely
            # We can't simply instantiate Paragraph(new_element, parent) easily without internal API
            # So we use a trick: we find it in the document (it's now part of the tree)
            # But searching is slow.
            
            # Better way: Modify the XML directly for the text, since we just cloned it
            # The cloned element has the OLD text runs. We need to clear/update them.
            
            # Clear all run texts in the XML
            # w:r is the run namespace
            for run_idx, run in enumerate(new_element.findall(qn('w:r'))):
                # Find the w:t (text) elements
                texts = run.findall(qn('w:t'))
                if run_idx == 0:
                    # Set the text of the first run to our new content
                    if texts:
                        texts[0].text = new_text
                    else:
                        # Create w:t if missing (unlikely in a clone of text para)
                        pass 
                    # Remove other text elements in this run if any
                    for t in texts[1:]:
                         run.remove(t)
                else:
                    # Remove subsequent runs
                    new_element.remove(run)
            
            # If the first run had no text element (unlikely), we might need to add it.
            # But assuming the source had text, the clone has text.
            
            # Update our cursor to this new paragraph (for the next iteration)
            # We don't fully need the proxy object if we continue cloning the XML
            current_para = type('obj', (object,), {'_element': new_element}) # Mock object with _element
            
            # Add to our map? 
            # Ideally yes, but we might not need it for this transaction. 
            self.para_map[new_id] = current_para # Sort of works if we need it later

    def _update_paragraph_text(self, paragraph, new_text: str):
        """
        Replace text while trying to preserve run-level formatting of the *start*.
        (Simple preservation strategy: keep the first run's style, clear others)
        """
        if not paragraph.runs:
            paragraph.add_run(new_text)
            return

        # Heuristic: The first run usually contains the "Paragraph Style" overrides (bold, etc)
        # We set the first run to the new text, and remove subsequent runs.
        # This destroys "intra-paragraph" formatting (e.g. one bold word in middle),
        # but preserves "paragraph-level" formatting (e.g. the whole thing is bold).
        
        # 1. Update first run
        paragraph.runs[0].text = new_text
        
        # 2. Remove subsequent runs
        # We iterate backwards to avoid index shifting issues, checking we don't delete run 0
        for i in range(len(paragraph.runs) - 1, 0, -1):
            # python-docx doesn't have a clean 'remove_run', so we clear text
            # OR we can remove the element from xml. Clearing text is safer/easier.
            paragraph.runs[i].text = "" 
            # If we really want to remove the xml element:
            # r = paragraph.runs[i]._element
            # r.getparent().remove(r)

    def _save(self) -> bytes:
        target_stream = io.BytesIO()
        self.doc.save(target_stream)
        target_stream.seek(0)
        return target_stream.read()
