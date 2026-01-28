import uuid
import io
from docx import Document
from docx.oxml.ns import qn

class IDNormalizer:
    """
    The 'Ingestion Gatekeeper'.
    Ensures every paragraph in a DOCX file has a stable, unique 'paraId'.
    This allows us to rely on these IDs for round-trip editing later.
    """
    
    @staticmethod
    def normalize_docx(file_content: bytes) -> bytes:
        """
        Parse the DOCX bytes, inject IDs where missing, and return new bytes.
        """
        try:
            doc = Document(io.BytesIO(file_content))
        except Exception as e:
            # If we can't parse it (e.g. PDF or bad file), returns original
            # Let the downstream parsers handle the error or treat as read-only
            return file_content

        modified = False
        existing_ids = set()

        # First pass: Collect existing IDs to ensure uniqueness
        for p in doc.paragraphs:
            # Access the underlying XML element
            # The attribute is w:paraId in the w namespace
            para_id = p._element.get(qn('w:paraId'))
            if para_id:
                existing_ids.add(para_id)

        # Second pass: Inject missing IDs
        for p in doc.paragraphs:
            para_id = p._element.get(qn('w:paraId'))
            
            if not para_id:
                # Generate a new 8-char hex ID (standard Word format is 8 hex chars)
                # But a UUID is fine too, though Word uses 8-char hex. 
                # Let's use a robust UUID to avoid any collision risk, 
                # even if it looks "non-native" to Word (Word tolerates it).
                new_id = str(uuid.uuid4())
                
                # Check collision (extremely unlikely with UUID4)
                while new_id in existing_ids:
                    new_id = str(uuid.uuid4())
                
                p._element.set(qn('w:paraId'), new_id)
                existing_ids.add(new_id)
                modified = True

        if not modified:
            return file_content

        # Save to bytes
        target_stream = io.BytesIO()
        doc.save(target_stream)
        target_stream.seek(0)
        return target_stream.read()
