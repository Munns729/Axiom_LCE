"""
Document parsing service for DOCX, PDF, and TXT files
"""
import io
from docx import Document
from pypdf import PdfReader
from typing import Tuple

class DocumentService:
    
    @staticmethod
    def extract_text_from_docx(file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(io.BytesIO(file_content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text = "\n\n".join(paragraphs)
            return text
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")
    
    @staticmethod
    def extract_text_from_pdf(file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            reader = PdfReader(pdf_file)
            
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
            
            text = "\n\n".join(text_parts)
            return text
        except Exception as e:
            raise ValueError(f"Error parsing PDF: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_content: bytes) -> str:
        """Extract text from TXT file"""
        try:
            return file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Try other common encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    return file_content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            raise ValueError("Could not decode text file with any common encoding")
    
    @staticmethod
    def extract_text(filename: str, content: bytes) -> Tuple[str, dict]:
        """
        Extract text and structure from file based on extension
        Returns: (text, tree_dict)
        """
        filename_lower = filename.lower()
        
        if filename_lower.endswith('.docx'):
            return DocumentService.parse_docx_structure(content)
        elif filename_lower.endswith('.pdf'):
            text = DocumentService.extract_text_from_pdf(content)
            return text, None
        elif filename_lower.endswith('.txt'):
            text = DocumentService.extract_text_from_txt(content)
            return text, None
        else:
            raise ValueError(
                f"Unsupported file type: {filename}. "
                "Supported types: .docx, .pdf, .txt"
            )
            
    @staticmethod
    def parse_docx_structure(content: bytes) -> Tuple[str, dict]:
        """
        Robustly parse DOCX into text and structure tree
        Uses logic ported from Spine for accurate clause detection
        Returns: Tuple[full_text, root_node_dict]
        """
        import re
        import uuid
        
        try:
            doc = Document(io.BytesIO(content))
        except Exception as e:
            raise ValueError(f"Error parsing DOCX: {str(e)}")

        paragraphs_text = []
        
        # Root node dict (matches ClauseNode schema)
        root = {
            "id": "root",
            "an_type": "document",
            "text_content": "ROOT",
            "children": [],
            "metadata": {}
        }
        
        current_article = None
        current_section = None
        
        for p in doc.paragraphs:
            text = p.text.strip()
            if not text:
                continue
                
            paragraphs_text.append(text)
            style_name = p.style.name
            
            # Safe ID strategy
            try:
                from docx.oxml.ns import qn
                if p._element is not None:
                    # Use standard qualified name lookup
                    para_id = p._element.get(qn('w:paraId'))
                    if not para_id:
                        # Fallback if Normalizer hasn't run (should verify logic)
                        para_id = str(id(p._element))
                else:
                    para_id = str(id(p))
            except Exception:
                para_id = str(uuid.uuid4())

            # --- Detection Logic (Robust) ---
            # Patterns for generic numbering detection
            pattern_section = r'^(\d+(\.\d+)*)\.?\s+(.*)' # 1.1 or 1.1.1 or 4
            pattern_point = r'^(\([a-z0-9]+\))\s+(.*)' # (a) or (1)
            
            # Base node structure matching schemas_ast.ClauseNode
            node = {
                "id": str(uuid.uuid4()),
                "text_content": text,
                "original_xml_id": para_id,
                "children": [],
                "an_type": "paragraph", # Default
                "an_num": None,
                "metadata": {}
            }
            
            # Priority 1: Headers (Explicit or Keyword)
            pattern_article = r'^(ARTICLE|SECTION|SCHEDULE|EXHIBIT)\s+([IVXLCDM0-9A-Z]+)' # Detect ARTICLE I or SECTION 1
            a_match = re.match(pattern_article, text, re.IGNORECASE)

            if style_name.startswith('Heading 1') or style_name == 'Title' or a_match:
                node["an_type"] = "article"
                if a_match:
                     node["an_num"] = f"{a_match.group(1)} {a_match.group(2)}"
                
                root["children"].append(node)
                current_article = node
                current_section = None 
                
            elif style_name.startswith('Heading 2'):
                match = re.match(pattern_section, text)
                node["an_type"] = "section"
                node["an_num"] = match.group(1) if match else None
                
                if current_article:
                    current_article["children"].append(node)
                else:
                    root["children"].append(node)
                current_section = node
            
            # Priority 2: Regex on Normal/List Text
            else:
                s_match = re.match(pattern_section, text)
                p_match = re.match(pattern_point, text)
                
                if s_match:
                     # Looks like section (1.1) but styled Normal
                     node["an_type"] = "section"
                     node["an_num"] = s_match.group(1)
                     
                     if current_article:
                         current_article["children"].append(node)
                     else:
                         root["children"].append(node)
                     current_section = node
                     
                elif p_match:
                    # Looks like point (a)
                    node["an_type"] = "point"
                    node["an_num"] = p_match.group(1)
                    
                    if current_section:
                        current_section["children"].append(node)
                    elif current_article:
                         current_article["children"].append(node)
                    else:
                        root["children"].append(node)
                        
                else:
                    # Standard Paragraph
                    node["an_type"] = "paragraph"
                    
                    if current_section:
                        current_section["children"].append(node)
                    elif current_article:
                        current_article["children"].append(node)
                    else:
                        root["children"].append(node)
            
            # --- Clause Classification ---
            node["clause_type"] = DocumentService._detect_clause_type(text)

        full_text = "\n\n".join(paragraphs_text)
        return full_text, root

    @staticmethod
    def _detect_clause_type(text: str) -> str:
        """Simple heuristic to classify clause type"""
        text_lower = text.lower()
        
        # Conditions
        if any(w in text_lower for w in ["if ", "unless", "provided that", "subject to", "condition"]):
            return "condition"
            
        # Obligations
        if any(w in text_lower for w in ["shall", "must", "agree to", "agrees to", "will"]):
            return "obligation"
            
        # Rights / Permissions
        if any(w in text_lower for w in ["may", "entitled to", "right to", "option to"]):
            return "right"
            
        # Representations
        if any(w in text_lower for w in ["represents", "warrants", "representation", "warranty"]):
             return "representation"
             
        # Definitions
        if any(w in text_lower for w in ["means", "defined as", "meaning"]):
            return "definition"
            
        return "general"

    @staticmethod
    def format_file_size(size_bytes: int) -> str:
        """Format file size in human-readable format"""
        for unit in ['B', 'KB', 'MB', 'GB']:
            if size_bytes < 1024.0:
                return f"{size_bytes:.1f} {unit}"
            size_bytes /= 1024.0
        return f"{size_bytes:.1f} TB"
