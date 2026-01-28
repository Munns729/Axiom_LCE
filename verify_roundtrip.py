import os
import sys
import json
import uuid
# Add backend to path so we can import services
sys.path.append(os.path.join(os.getcwd(), 'backend'))

from services.id_normalizer import IDNormalizer
from services.composer_service import ComposerService
# Import the generator from spine
sys.path.append(os.path.join(os.getcwd(), 'spine'))
from generate_complex_corpus import create_series_a_document
from docx import Document
from docx.oxml.ns import qn

def verify_phase_a():
    print("=== STARTING PHASE A VERIFICATION ===")
    
    # 1. Generate a fresh Complex Document
    print("[1] Generating Complex Series A Document...")
    raw_path = create_series_a_document()
    with open(raw_path, "rb") as f:
        original_bytes = f.read()
    print(f"    Created: {raw_path} ({len(original_bytes)} bytes)")

    # 2. Run ID Normalizer (Ingestion)
    print("[2] Running ID Normalizer...")
    normalized_bytes = IDNormalizer.normalize_docx(original_bytes)
    print(f"    Normalized Size: {len(normalized_bytes)} bytes")
    
    # Verify IDs exist and capture properties
    doc = Document(io.BytesIO(normalized_bytes))
    target_node_id = None
    target_text = "Bad Leaver" # Indented paragraph
    original_indent = None
    
    for p in doc.paragraphs:
        pid = p._element.get(qn('w:paraId'))
        if target_text in p.text:
            target_node_id = pid
            original_indent = p.paragraph_format.left_indent
            print(f"    Found Target '{target_text}' -> ID: {pid}")
            print(f"    Original Indent: {original_indent} (Should be ~457200 for 36pt)")
            break

    if not target_node_id:
        print("FAILED: Could not find target paragraph.")
        return

    # 3. Initialize Composer (The Bridge)
    print("[3] Initializing ComposerService...")
    composer = ComposerService(normalized_bytes)

    # 4. Execute Surgical Split (Phase B)
    print(f"[4] Executing SPLIT on Node {target_node_id}...")
    part_a = "3.1 [PART A] All Shares held by the Founders shall be subject to reverse vesting."
    part_b = "3.2 [PART B - NEW] The Vesting Period shall be 48 months."
    
    ops = [
        {
            "type": "split", 
            "id": target_node_id, 
            "parts": [part_a, part_b]
        }
    ]
    
    final_bytes = composer.apply_operations(ops)
    
    # 5. Save and Verify Result
    output_path = os.path.join(os.getcwd(), 'series_a_roundtrip.docx')
    with open(output_path, "wb") as f:
        f.write(final_bytes)
        
    print(f"[5] Saved Result: {output_path}")
    
    # Re-read to prove change
    check_doc = Document(output_path)
    found_a = False
    found_b = False
    
    for p in check_doc.paragraphs:
        if "[PART A]" in p.text:
            found_a = True
            print(f"    SUCCESS: Found Part A: '{p.text}'")
            # Verify ID persisted
            if p._element.get(qn('w:paraId')) == target_node_id:
                print("    SUCCESS: Part A kept original ID.")
            else:
                 print(f"    WARNING: Part A Lost ID!")
                 
        if "[PART B - NEW]" in p.text:
            found_b = True
            print(f"    SUCCESS: Found Part B: '{p.text}'")
            # Verify it has a NEW ID
            new_id = p._element.get(qn('w:paraId'))
            if new_id and new_id != target_node_id:
                print(f"    SUCCESS: Part B has new ID: {new_id}")
            else:
                print(f"    FAILURE: Part B has bad ID: {new_id}")
            
            # Verify formatting Inheritance
            # We assume it inherited the indent from the Bad Leaver clause (36pt)
            new_indent = p.paragraph_format.left_indent
            if new_indent == original_indent:
                 print("    SUCCESS: Part B inherited indentation (Style Cloning Worked)!")
            else:
                 print(f"    FAILURE: Part B lost formatting! {original_indent} -> {new_indent}")

    if found_a and found_b:
        print("\n=== VERIFICATION PASSED: SPLIT COMPLETE ===")
    else:
        print("\n=== VERIFICATION FAILED: MISSING PARTS ===")

import io
if __name__ == "__main__":
    verify_phase_a()
